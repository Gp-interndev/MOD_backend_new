from docx import Document

template_path = r"D:\Monarch_Mod\backend\MOD 3.docx"
docmonarch = Document(template_path)

# Ensure the document has at least two tables
if len(docmonarch.tables) > 1:
    # Assign tables before modifying or iterating
    table = docmonarch.tables[0]  # First table
    table1 = docmonarch.tables[1]  # Second table

    # Print all tables
    for table_index, table_obj in enumerate(docmonarch.tables):
        print(f"Table {table_index + 1}:\n")
        for row in table_obj.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print("\t".join(row_text))  # Print row text in a tab-separated format
        print("\n" + "=" * 50 + "\n")  # Separator between tables

    # Now you can safely use `table` and `table1` separately
    print("Table 1 and Table 2 have been assigned successfully.")

else:
    print("Document does not contain enough tables.")
