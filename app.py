from flask import Flask, request, jsonify, render_template, send_file
import psycopg2
from datetime import datetime
from flask_cors import CORS
import pandas as pd
import json
import sys
import math
import folium
from folium.plugins import MeasureControl
from geopy.distance import geodesic
from pyproj import Proj, transform
import re
import geopandas as gpd
from shapely.geometry import Polygon
from shapely.ops import nearest_points
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import requests
import os
# import comtypes.client
import logging
# import pythoncom  
# import pypandoc
from docx2pdf import convert
import platform
from docx.enum.text import WD_ALIGN_PARAGRAPH
# import geopandas as gpd



app = Flask(__name__)
CORS(app) 

# Database Configuration    
DB_HOST = "iwmsgis.pmc.gov.in"
DB_NAME = "MOD"
DB_USER = "postgres"
DB_PASS = "pmc992101"
DB_PORT = "5432" 

def get_db_connection():
    """Establish connection to PostgreSQL"""
    return psycopg2.connect(
        host=DB_HOST,
        database=DB_NAME,
        user=DB_USER,
        password=DB_PASS
    )


# Database connection For Update_csv api
# DB_NAME = "mod"
# DB_USER = "postgres"
# DB_PASSWORD = "Mojani@992101"
# DB_HOST = "info.dpzoning.com"  #info.dpzoning.com
# DB_PORT = "5432"  




# Login Route
@app.route('/admin_login', methods=['POST'])
def admin_login():
    conn = None
    cursor = None
    try:
        data = request.json
        username = data.get("username")
        password = data.get("password")

        if not username or not password:
            return jsonify({"error": "Username and password are required"}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # Fetch user from DB
        cursor.execute("SELECT password FROM admin_users WHERE username = %s", (username,))
        user = cursor.fetchone()

        if not user:
            return jsonify({"error": "Invalid credentials"}), 401

        stored_password = user[0]

        # Compare password directly (plain text comparison)
        if password == stored_password:
            return jsonify({"message": "Login successful"}), 200
        else:
            return jsonify({"error": "Invalid credentials"}), 401

    except psycopg2.Error as e:
        return jsonify({"error": str(e)}), 500
    finally:
        # Close cursor and connection only if they were initialized
        if cursor:
            cursor.close()
        if conn:
            conn.close()



@app.route('/save_user', methods=['POST'])
def save_user():
    """API endpoint to insert user data into PostgreSQL"""
    try:
        data = request.json  # Get JSON input from request
        
        # Extract data from JSON payload
        name = data.get("name")
        mobilenumber = data.get("mobilenumber")
        nameoncertificate = data.get("nameoncertificate")
        gstnumber = data.get("gstnumber") if data.get("gstnumber") else None
        pannumber = data.get("pannumber") if data.get("pannumber") else None
        # siteadress = data.get("siteadress")
        gutnumber = data.get("gutnumber") 
        district = data.get("district")
        taluka = data.get("taluka")
        village = data.get("village")
        pincode = data.get("pincode") if data.get("pincode") else None
        correspondanceadress = data.get("correspondanceadress")
        # outwardnumber = data.get("outwardnumber")
        date = datetime.now()  # Store current timestamp

        if not all([name, nameoncertificate, gutnumber, district, taluka, village]):
            return jsonify({"error": "Missing required fields"}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # SQL Query to Insert Data
        insert_query = """
        INSERT INTO public.userdata 
        (name, mobilenumber, nameoncertificate, gstnumber, pannumber, gutnumber, 
         district, taluka, village, pincode, correspondanceadress,  date) 
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        RETURNING outwardnumber
        """
        cursor.execute(insert_query, (name, mobilenumber, nameoncertificate, gstnumber, pannumber, 
                                       gutnumber, district, taluka, village, pincode, 
                                      correspondanceadress,  date))

        outwardnumber = cursor.fetchone()[0]
        conn.commit()

        cursor.close()
        conn.close()

        return jsonify({"message": "User data saved successfully!",
                        "outwardnumber": outwardnumber})

    except psycopg2.Error as e:
        return jsonify({"error": str(e)}), 500
    

@app.route('/get_user/<string:outwardnumber>', methods=['GET'])
def get_user_by_outwardnumber(outwardnumber):
    """API endpoint to retrieve a single user by outwardnumber"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Query with explicit column selection to avoid incorrect mapping
        query = """SELECT outwardnumber, name, mobilenumber, nameoncertificate, gstnumber, 
                   pannumber, gutnumber, district, taluka, village, 
                   pincode, correspondanceadress, date FROM userdata WHERE outwardnumber = %s"""
        cursor.execute(query, (outwardnumber,))
        user = cursor.fetchone()

        cursor.close()
        conn.close()

        if user:
            # Define column names explicitly in the correct order
            columns = ["outwardnumber", "name", "mobilenumber", "nameoncertificate", "gstnumber", 
                       "pannumber", "gutnumber", "district", "taluka", "village", 
                       "pincode", "correspondanceadress", "date"]
 
            user_data = dict(zip(columns, user))  # Convert tuple to dictionary

            return jsonify({"user": user_data}), 200
        else:
            return jsonify({"message": "User not found"}), 404

    except psycopg2.Error as e:
        return jsonify({"error": str(e)}), 500



# API for the Coordinates Data
# Define UTM projection for Zone 43N
utm_proj = Proj(proj="utm", zone=43, datum="WGS84", south=False)
wgs84_proj = Proj(proj="latlong", datum="WGS84")

# Function to convert decimal degrees to DMS
def decimal_to_dms(decimal_degree):
    degrees = int(decimal_degree)
    minutes = int((abs(decimal_degree) - abs(degrees)) * 60)
    seconds = (abs(decimal_degree) - abs(degrees) - minutes / 60) * 3600
    return degrees, minutes, seconds

# Function to calculate distance using the Haversine formula
def haversine(lat1, lon1, lat2, lon2):
    R = 6371.0  # Radius of Earth in kilometers
    lat1, lon1, lat2, lon2 = map(math.radians, [lat1, lon1, lat2, lon2])
    dlat = lat2 - lat1
    dlon = lon2 - lon1
    a = math.sin(dlat / 2) ** 2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon / 2) ** 2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    return R * c

# def calculate_distance(lat1, lon1, lat2, lon2):
#     # Create Point objects
#     point1 = Point(lon1, lat1)
#     point2 = Point(lon2, lat2)
    
#     # Create a GeoDataFrame
#     gdf = gpd.GeoDataFrame(geometry=[point1, point2], crs="EPSG:4326")
    
#     # Calculate the distance
#     distance = gdf.distance(gdf.shift()).iloc[1]  # Distance between the two points
#     return distance / 1000  # Convert meters to kilometers




def map_sattelite(coords, points_with_labels,nearest_points_list, output_map="static/map.html"):
    """
    Create a folium map with a polygon, labeled points, and an export-to-PDF button on Google Satellite imagery.
    
    Args:
        coords (list): List of (latitude, longitude) tuples for the polygon.
        points_with_labels (list): List of tuples [(lat, lon, label), ...] for points with labels.
        output_map (str): Path to save the output HTML map.
    
    Returns:
        str: Path to the saved HTML map.
    """
    swapped_coords = [(lat,lon) for lon, lat  in coords]
    m = folium.Map(
        # location=[coords[0][0], coords[0][1]],
        location = swapped_coords[0],
        zoom_start=10,
        tiles=None  # Disable default tiles
    )
    m.add_child(MeasureControl())

    folium.TileLayer(
        tiles="https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png",
        attr="OpenStreetMap",
        name="OpenStreetMap",
        overlay=False,
        control=True  # Allow users to toggle this layer
    ).add_to(m)
    
    # Add Google Satellite Tiles
    folium.TileLayer(
        tiles="https://mt1.google.com/vt/lyrs=s&x={x}&y={y}&z={z}",
        attr="Google Satellite",
        name="Google Satellite",
        overlay=False,
        control=True
    ).add_to(m) 
    
    polygon = folium.Polygon(
        locations=swapped_coords,  # List of (latitude, longitude) tuples
        color="red",
        weight=3,
        fill=True,
        fill_color="cyan",
        fill_opacity=0.4,
        popup="Polygon Area"
    ).add_to(m)

    Aviation_boundary = folium.WmsTileLayer(
        url="https://iwmsgis.pmc.gov.in/geoserver/wms?",
        name="Aviation Boundaries",
        layers="MOD:Aviation_Boundary",
        fmt="image/png",
        transparent=True,
        overlay=True,
        control=True
    ).add_to(m)

    
    Aviation_zone =folium.WmsTileLayer(
        url="https://iwmsgis.pmc.gov.in/geoserver/wms?",
        name="Aviation Zone",
        layers="MOD:Aviation_data",
        fmt="image/png",
        transparent=True,
        overlay=True,
        opacity = 0.5,
        control=True
    ).add_to(m)

   
    for point_pair in nearest_points_list:
    # Each point_pair is a tuple of two points
        point1 = point_pair[0]  
        point2 = point_pair[1] 

        # Calculate the distance between the two points using geodesic (this calculates the great-circle distance)
        line_length = geodesic(point1, point2).kilometers  # Distance in kilometers

        mid_point_lat = (point1[0] + point2[0]) / 2
        mid_point_lon = (point1[1] + point2[1]) / 2
        popup_message = f"Distance: {line_length:.2f} km"  # Format the distance to two decimal places

        # Add the PolyLine to the map with the popup showing the distance
        folium.PolyLine(
            locations=[point1, point2],  # Coordinates of the points to draw a line between
            color="yellow",  # Color for the line
            weight=1,  # Line thickness
        ).add_to(m).add_child(folium.Popup(popup_message))

        folium.Marker(
        location=[mid_point_lat, mid_point_lon],  # Midpoint of the line
        icon=folium.DivIcon(
            icon_size=(150, 36),  # Size of the label
            icon_anchor=(7, 20),  # Position of the label
            html=f'<div style="font-size: 16px; font-weight: bold; color: yellow;">{line_length:.2f} km</div>'  # Label style
        ),
    ).add_to(m)


    for lat, lon, label in points_with_labels:
        folium.CircleMarker(
            location=(lat, lon),
            radius=3,  # Small dot size
            color="blue",
            fill=True,
            fill_color="blue",
            fill_opacity=0.5,
            popup=f"{label}",  # Add label as a popup
        ).add_to(m)

        folium.Marker(
        location=[lat, lon], 
        icon=folium.DivIcon(
            icon_size=(150, 36),  # Size of the label
            icon_anchor=(7, 20),  # Position of the label
            html=f'<div style="font-size: 12px; font-weight: bold; color: yellow;">{label}</div>'  # Label style
        ),
    ).add_to(m)

    m.fit_bounds(polygon.get_bounds()) 
    # Add a custom button to export to PDF
    pdf_button = """
      <script src="https://cdnjs.cloudflare.com/ajax/libs/html2canvas/1.4.1/html2canvas.min.js" integrity="sha512-BNaRQnYJYiPSqHHDb58B0yaPfCu+Wgds8Gp/gU33kqBtgNS4tSPHuGibyoeqMV/TJlSKda6FXzoEyYGjTe+vXA==" crossorigin="anonymous" referrerpolicy="no-referrer"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/jspdf/2.5.1/jspdf.umd.min.js"></script>
                <div style="position: fixed; 
                            bottom: 50px; left: 50px; width: 150px; height: 30px; 
                            z-index: 1000;">
                    <button onclick="exportToPDF()" style="width: 150px; height: 30px; background-color: #4CAF50; color: white; border: none; border-radius: 5px;">
                        Export to PDF
                    </button>
                </div>
            
                <script>
                function exportToPDF() {
    try {
        const { jsPDF } = window.jspdf;
        const doc = new jsPDF();

        // Select the map container dynamically
        const mapContainer = document.querySelector('.folium-map');

        const originalScrollX = window.scrollX;
        const originalScrollY = window.scrollY;

        html2canvas(mapContainer, {
            scale: 2, // Scale for high resolution
            useCORS: true, // Handle cross-origin images
            scrollX: originalScrollX, // Maintain original horizontal scroll position
            scrollY: originalScrollY, // Maintain original vertical scroll position
        }).then(function (canvas) {
            const imgData = canvas.toDataURL('image/png');
            const pdfWidth = 180; // Maximum width for PDF
            const aspectRatio = canvas.width / canvas.height;
            const imgHeight = pdfWidth / aspectRatio; // Maintain aspect ratio

            // Center the map image in the PDF
            const pageWidth = doc.internal.pageSize.getWidth();
            const centerX = (pageWidth - pdfWidth) / 2;

            // Add the image to the PDF
            doc.addImage(imgData, 'PNG', centerX, 10, pdfWidth, imgHeight);

            // Save the generated PDF
            const fileName = `map_export.pdf`;
            doc.save(fileName);
        });
    } catch (error) {
        console.error('Error generating PDF:', error);
        alert('Failed to generate PDF. Check console for details.');
    }
}


    </script>
    """
    m.get_root().html.add_child(folium.Element(pdf_button))

    folium.LayerControl().add_to(m)
    m.save(output_map)
    return output_map

def convert_to_wgs84(x, y):
    lon, lat = transform(utm_proj, wgs84_proj, x, y)
    return lat, lon


def calculate_boundaryDistance(coords):

    wmsUrlNDALohgaonBOundary = "http://iwmsgis.pmc.gov.in/geoserver/ows?service=WFS&version=2.0.0&request=GetFeature&typeName=MOD:Aviation_Boundary&outputFormat=application/json"
    geoserver_layer = gpd.read_file(wmsUrlNDALohgaonBOundary)
    geoserver_layer = geoserver_layer.to_crs(epsg=32643) 
    polygon_nda = geoserver_layer[geoserver_layer["Aviation_N"] == "NDA"]
    polygon_lohgaon = geoserver_layer[geoserver_layer["Aviation_N"] == "Lohagaon"]
    

    # Create polygon from input coordinates
    polygon_layout = gpd.GeoDataFrame(
        {'geometry': [Polygon(coords)]},
        crs="EPSG:32643"  # Original CRS for input coordinates (WGS84)
    ).geometry.iloc[0]
   
    polygon_nda = polygon_nda.to_crs(epsg=32643).geometry.iloc[0]
    polygon_lohgaon = polygon_lohgaon.to_crs(epsg=32643).geometry.iloc[0]

    # calculate nearest point
    nearest_nda_point = nearest_points(polygon_layout, polygon_nda)[1]
    nearest_lohgaon_point = nearest_points(polygon_layout, polygon_lohgaon)[1]
    polygon_layout_NDA = nearest_points(polygon_nda,polygon_layout)[1]
    polygon_layout_Lohagaon = nearest_points(polygon_lohgaon,polygon_layout)[1]

    nearest_nda_point_wgs84 = convert_to_wgs84(nearest_nda_point.x, nearest_nda_point.y)
    nearest_lohgaon_point_wgs84 = convert_to_wgs84(nearest_lohgaon_point.x, nearest_lohgaon_point.y)
    nearest_polygon_layout_NDA_wgs84 = convert_to_wgs84(polygon_layout_NDA.x,polygon_layout_NDA.y)
    nearest_polygon_layout_Lohagaon_wgs84 = convert_to_wgs84(polygon_layout_Lohagaon.x,polygon_layout_Lohagaon.y)

     # calculate distance point
    distance_meters_nda = polygon_nda.distance(polygon_layout)
    distance_meters_lohgaon = polygon_lohgaon.distance(polygon_layout)
    distance_km_nda = distance_meters_nda / 1000

    # print(distance_km_nda,"ooooooooooooooooooooooooooooooooooooooooooooooooooooo")
    distance_km_lohgaon = distance_meters_lohgaon / 1000
    
    # Return the distances in a dictionary
    mindistance = {
        "NDAboundaryMinDistance": distance_km_nda,
        "LohgaonBoundaryMinDistance": distance_km_lohgaon
    }
    nearest_points_list = [
        [nearest_nda_point_wgs84, nearest_polygon_layout_NDA_wgs84],  # (lat, lon) format for folium
        [nearest_lohgaon_point_wgs84, nearest_polygon_layout_Lohagaon_wgs84]
    ] 
    return mindistance,nearest_points_list

# Route to handle CSV file upload and processing






ALLOWED_EXTENSIONS = {'csv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/process_csv', methods=['POST'])
def process_csv():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file and allowed_file(file.filename):
        try:
            filename = secure_filename(file.filename)
            data = pd.read_csv(file,header=None)

            print("Columns in CSV:", data.columns)
            
            if data.shape[1] < 3:
                return jsonify({"error": "CSV file must contain at least 3 columns (P name, UTM x, UTM y)"}), 400

            decimal_degrees = []
            fpoints = []
            utmpoints = []
            fpointswithlabel = []
            reference_points = {
                "NDA": {"utm_x": 371129.923, "utm_y": 2042927.865},
                "loh": {"utm_x": 385999.526, "utm_y": 2055079.640},
            }

            # Create transformer object for coordinate conversion
            utm = Proj('epsg:32643')  # UTM zone 43N
            wgs84 = Proj('epsg:4326')  # WGS84

            for ref_name, ref_coords in reference_points.items():
                # Convert reference points from UTM to WGS84
                lon, lat = transform(utm, wgs84, ref_coords["utm_x"], ref_coords["utm_y"])
                reference_points[ref_name]["latitude"] = float(lat)
                reference_points[ref_name]["longitude"] = float(lon)

            # Iterate through data rows using itertuples()
            for row in data.itertuples():
                p_name = str(row[1])  # First column
                x = float(row[2])     # Second column (UTM X)
                y = float(row[3])     # Third column (UTM Y)
                elevation_val = float(row[4]) if len(row) > 4 else None  # Fourth column if exists

                # Convert from UTM to WGS84
                lon, lat = transform(utm, wgs84, x, y)
                lat = float(lat)
                lon = float(lon)
                
                lat_dms = decimal_to_dms(lat)
                lon_dms = decimal_to_dms(lon)
                utmpoint = (x, y)
                utmpoints.append(utmpoint)
                
                distances = {}
                for ref_name, ref_coords in reference_points.items():
                    distances[ref_name] =(haversine(lat, lon, ref_coords["latitude"], ref_coords["longitude"]))

                if isinstance(p_name, str) and re.match(r"^\s*[Pp]", p_name):
                    points = (float(lat), float(lon))
                    pointslabel = (float(lat), float(lon), str(p_name))
                    fpoints.append(points)
                    fpointswithlabel.append(pointslabel)

                decimal_degrees.append({
                    "P_name": p_name,
                    "latitude": lat,
                    "longitude": lon,
                    "Height": elevation_val,
                    "longitude_dms": f"{lat_dms[0]}°{lat_dms[1]}'{lat_dms[2]:.2f}\"",
                    "latitude_dms": f"{lon_dms[0]}°{lon_dms[1]}'{lon_dms[2]:.2f}\"",
                    "distances_to_reference_points_km": distances,
                })

            boundary_distances, nearest_points_list = calculate_boundaryDistance(utmpoints)
            
            # Convert boundary distances to float
            boundary_distances = {
                "NDAboundaryMinDistance": float(boundary_distances["NDAboundaryMinDistance"]),
                "LohgaonBoundaryMinDistance": float(boundary_distances["LohgaonBoundaryMinDistance"])
            }
            print(fpoints, fpointswithlabel, nearest_points_list)
            map_sattelite(fpoints, fpointswithlabel, nearest_points_list)

            result = {
                "decimal_degrees": decimal_degrees,
                "boundary_distances": boundary_distances,
                "Height of plot from sea surface":"573.7 + 15.35 = 589.05 M",
                "Height of building from sea surface":"573.4 + 15.35 = 588.05 M"
            }

            return jsonify(result), 200

        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"Error details: {error_details}")
            return jsonify({"error": str(e), "details": error_details}), 500

    return jsonify({"error": "Invalid file format. Only CSV files are allowed."}), 400




@app.route('/update_csv', methods=['POST'])
def update_csv():
    # Get outward number from the form data
    outwardnumber = request.form.get('outwardNumber')
    # outwardnumber = '1069'
    if not outwardnumber:
        return jsonify({"error": "Outward number is required"}), 400
    
    file = request.files.get('file')  
    if file and allowed_file(file.filename):
        try:
            # First, fetch user data using the outward number
            conn = psycopg2.connect(dbname=DB_NAME, user=DB_USER, password=DB_PASS, host=DB_HOST, port=DB_PORT)
            cur = conn.cursor()
            
           
            user_query = """
                SELECT name, district, taluka, village, date, correspondanceadress, gutnumber
                FROM userdata WHERE outwardnumber = %s
            """
            cur.execute(user_query, (outwardnumber,))
            user_data = cur.fetchone()
            
            if not user_data:
                return jsonify({"error": f"No user found with outward number: {outwardnumber}"}), 404
                
            # Extract user data
            user_name, district, taluka, village, date, address, gut = user_data
            
            # CSV processing
            filename = secure_filename(file.filename)
            data = pd.read_csv(file,header=None)
            pointplotCoordiantes, pointbuildingCoordiantes = [], []
            plotCoordiantes, buildingCoordiantes = [], []

            print("Columns in CSV:", data.columns)
            print(f"Using outward number: {outwardnumber}")

            if data.shape[1] < 3:
                return jsonify({"error": "CSV file must contain at least 3 columns (P name, UTM x, UTM y)"}), 400

            
            reference_points = {
                "NDA": {"utm_x": 371129.923, "utm_y": 2042927.865},
                "loh": {"utm_x": 385999.526, "utm_y": 2055079.640},
            }

            # Create transformer object for coordinate conversion
            utm = Proj('epsg:32643')  
            wgs84 = Proj('epsg:4326') 

            
            for ref_name, ref_coords in reference_points.items():
                lon, lat = transform(utm, wgs84, ref_coords["utm_x"], ref_coords["utm_y"])
                reference_points[ref_name]["latitude"] = float(lat)
                reference_points[ref_name]["longitude"] = float(lon)

            print(reference_points,"ppppopopiouiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiiii")
            for index, row in data.iterrows():
                labelName, utmXcoordiantes, utmYcoordinates, Zcoordinates = row[0], row[1], row[2], row[3]
                print(labelName, utmXcoordiantes, utmYcoordinates, Zcoordinates)
                pattern = r"^\s*[Pp]"

                
                lon, lat = transform(utm, wgs84, float(utmXcoordiantes), float(utmYcoordinates))
                lat, lon = float(lat), float(lon)
                
            
                nda_distance = haversine(lat, lon, reference_points["NDA"]["latitude"], reference_points["NDA"]["longitude"])
                print(nda_distance,"VVVVVVVVVVVVVVVVVVVVVVVVVVVVVVVVVVVVV")
                loh_distance = haversine(lat, lon, reference_points["loh"]["latitude"], reference_points["loh"]["longitude"])

                # Check if the label name starts with "P" or "p" (with optional spaces)
                if re.match(pattern, labelName):
                    print("------------------------------------------------------")
                    pointplotCoordiantes.append((labelName, utmXcoordiantes, utmYcoordinates, Zcoordinates, nda_distance, loh_distance))
                    plotCoordiantes.append((utmXcoordiantes, utmYcoordinates, Zcoordinates))
                else:
                    pointbuildingCoordiantes.append((labelName, utmXcoordiantes, utmYcoordinates, Zcoordinates, nda_distance, loh_distance))
                    buildingCoordiantes.append((utmXcoordiantes, utmYcoordinates, Zcoordinates))


            if plotCoordiantes and plotCoordiantes[0] != plotCoordiantes[-1]:
                plotCoordiantes.append(plotCoordiantes[0])

            if buildingCoordiantes and buildingCoordiantes[0] != buildingCoordiantes[-1]:
                buildingCoordiantes.append(buildingCoordiantes[0])

            # Create the polygon for plotCoordiantes
            plot_polygon_wkt = "POLYGONZ((" + ", ".join(f"{x} {y} {z}" for x, y, z in plotCoordiantes) + "))"
            
            # Create the polygon for buildingCoordiantes
            building_polygon_wkt = "POLYGONZ((" + ", ".join(f"{x} {y} {z}" for x, y, z in buildingCoordiantes) + "))"

            # Insert building coordinates if available
            if buildingCoordiantes:
                for name, x, y, z, nda_dist, loh_dist in pointbuildingCoordiantes:
                    point_wkt = f"POINTZ({x} {y} {z})"
                    # Use quoted column names to preserve case and include user data
                    cur.execute("""
                        INSERT INTO points (pointname, geom, outward, typeofsite, "Distance_from_NDA", "Distance_from_lohgaon",
                                            name, districtname, talukaname, villagename, date, address, gut, "Height_AMSL") 
                        VALUES (%s, ST_GeomFromText(%s, 32643), %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s);
                    """, (name, point_wkt, outwardnumber, "building", nda_dist, loh_dist,
                        user_name, district, taluka, village, date, address, gut, z))

                # Insert the building polygon into the `mod` table
                cur.execute("INSERT INTO mod (geom, outward, typeofsite) VALUES (ST_GeomFromText(%s, 32643), %s, %s);", 
                            (building_polygon_wkt, outwardnumber, "building"))

            # Insert plot coordinates 
            if plotCoordiantes:
                for name, x, y, z, nda_dist, loh_dist in pointplotCoordiantes:
                    point_wkt = f"POINTZ({x} {y} {z})"
                    
                    cur.execute("""
                        INSERT INTO points (pointname, geom, outward, typeofsite, "Distance_from_NDA", "Distance_from_lohgaon",
                                            name, districtname, talukaname, villagename, date, address, gut, "Height_AMSL") 
                        VALUES (%s, ST_GeomFromText(%s, 32643), %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s);
                    """, (name, point_wkt, outwardnumber, "plot", nda_dist, loh_dist,
                        user_name, district, taluka, village, date, address, gut, z))

                # Insert the plot polygon into the `mod` table
                cur.execute("INSERT INTO mod (geom, outward, typeofsite) VALUES (ST_GeomFromText(%s, 32643), %s, %s);", 
                            (plot_polygon_wkt, outwardnumber, "plot"))

           
            conn.commit()

            
            cur.close()
            conn.close()

            return jsonify({
                "message": "CSV data processed and inserted successfully!",
                "distances_added": True,
                "user_data_added": True
            })

        except Exception as e:
            print(f"Error occurred: {e}")
            import traceback
            error_details = traceback.format_exc()
            print(f"Error details: {error_details}")
            return jsonify({"error": f"An error occurred: {str(e)}", "details": error_details}), 500

    else:
        return jsonify({"error": "No file or invalid file format."}), 400


@app.route('/get_aviation_data/<string:outwardnumber>', methods=['GET'])
def get_aviation_data_and_geometry(outwardnumber):
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Step 1: Get the latest geometry for this outward number
        cur.execute("""
            SELECT geom FROM mod 
            WHERE outward = %s
            ORDER BY id DESC
            LIMIT 1;
        """, (outwardnumber,))

        result = cur.fetchone()

        if not result:
            return jsonify({"error": "No geometry found for this outward number"}), 404

        geometry = result[0]

        # Step 2: Query aviation data that intersects with the geometry
        cur.execute("""
            SELECT zone, elevation
            FROM "Aviation_data"
            WHERE ST_Intersects(geom, ST_Transform(ST_SetSRID(%s::geometry, 32643), 4326))
            LIMIT 1;
        """, (geometry,))

        aviation_data = cur.fetchone()

        # Add debug logging
        if aviation_data:
            print(f"Found aviation data: {aviation_data}")
        else:
            print(f"No aviation data found that intersects with this geometry")
            
            # Additional debugging query to check if Aviation_data table has any records
            cur.execute("SELECT COUNT(*) FROM Aviation_data")
            count = cur.fetchone()[0]
            print(f"Total records in Aviation_data: {count}")

        cur.close()
        conn.close()

        if not aviation_data:
            return jsonify({
                "geometry": geometry,
                "aviation_data": None,
                "message": "No aviation data found for this location"
            }), 200

        # Return both geometry and aviation data
        return jsonify({
            "geometry": geometry,
            "aviation_data": {
                "zone": aviation_data[0],
                "elevation": aviation_data[1]
            }
        }), 200

    except Exception as e:
        import traceback
        print(f"Error in get_aviation_data_and_geometry: {e}")
        print(traceback.format_exc())
        return jsonify({"error": str(e)}), 500



# @app.route('/get_aviation_data/<string:outwardnumber>', methods=['GET'])
# def get_aviation_data(outwardnumber):
#     try:
#         print(f"Fetching aviation data for outward number: {outwardnumber}")
#         conn = get_db_connection()
#         cur = conn.cursor()

#         # Get plot geometry for the given outward number
#         cur.execute("""
#             SELECT geom FROM mod 
#             WHERE outward = %s AND typeofsite = 'plot'
#         """, (outwardnumber,))
        
#         plot_geoms = cur.fetchall()  # Fetch all plots for the outward number

#         # Get building geometry for the given outward number
#         cur.execute("""
#             SELECT geom FROM mod 
#             WHERE outward = %s AND typeofsite = 'building'
#         """, (outwardnumber,))
        
#         building_geoms = cur.fetchall()  # Fetch all buildings for the outward number

#         if not plot_geoms and not building_geoms:
#             print(f"No plot or building geometry found for outward number: {outwardnumber}")
#             return jsonify({"error": "No plot or building geometry found for this outward number"}), 404

#         # Combine both plot and building geometries (if necessary)
#         all_geoms = plot_geoms + building_geoms
        
#         print(f"Found {len(all_geoms)} geometries for outward number: {outwardnumber}")

#         # Query aviation data that intersects with any of the geometries (plot or building)
#         aviation_data = None
#         for geom in all_geoms:
#             cur.execute("""
#                 SELECT zone, elevation
#                 FROM "Aviation_data"
#                 WHERE ST_Intersects(geom, ST_Transform(ST_SetSRID(%s::geometry, 32643), 4326))
#                 LIMIT 1;
#             """, (geom[0],))
            
#             aviation_data = cur.fetchone()
#             if aviation_data:
#                 break  # Exit the loop once aviation data is found
        
#         # Add debug logging
#         if aviation_data:
#             print(f"Found aviation data: {aviation_data}")
#         else:
#             print(f"No aviation data found that intersects with this plot/building")

#             # Additional debugging query to check if Aviation_data table has any records
#             cur.execute("SELECT COUNT(*) FROM Aviation_data")
#             count = cur.fetchone()[0]
#             print(f"Total records in Aviation_data: {count}")
        
#         cur.close()
#         conn.close()
        
#         if not aviation_data:
#             return jsonify({"error": "No aviation data found for this location"}), 404
        
#         # Return the actual aviation data
#         return jsonify({
#             "zone": aviation_data[0],
#             "elevation": aviation_data[1]
#         }), 200
        
#     except Exception as e:
#         import traceback
#         print(f"Error in get_aviation_data: {e}")
#         print(traceback.format_exc())
#         return jsonify({"error": str(e)}), 500



# Set up logging

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def set_table_borders(table):
    tbl = table._element
    tbl_pr = tbl.find(qn("w:tblPr"))
    
    # Ensure tblPr exists
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    
    # Set the table borders
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "5")  # Border size
        border.set(qn("w:space"), "0")  # Border space
        border.set(qn("w:color"), "000000")  # Border color
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

    # Prevent table rows from splitting across pages
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tbl_pr.append(cant_split)
    
    # Keep the tblLook element too
    tbl_look = OxmlElement("w:tblLook")
    tbl_look.set(qn("w:val"), "04A0")
    tbl_pr.append(tbl_look)

def set_paragraph_format(paragraph):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5  # Set line spacing
    paragraph_format.space_after = Pt(6)  # Space after paragraph
    paragraph_format.space_before = Pt(6)  # Space before paragraph

def set_cell_alignment(cell, vertical="center", horizontal="center"):
    """
    Set both vertical and horizontal alignment of the cell.
    vertical can be "top", "center", or "bottom".
    horizontal can be "left", "center", or "right".
    """
    # Get the cell's XML element
    tc = cell._element
    
    # Ensure the cell has a <w:tcPr> element
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc.insert(0, tc_pr)
    
    # Create <w:vAlign> element and set vertical alignment
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), vertical)
    tc_pr.append(v_align)

    # Set horizontal alignment
    for paragraph in cell.paragraphs:
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(horizontal, WD_ALIGN_PARAGRAPH.LEFT)


def adjust_table_cell_alignments(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_alignment(cell, vertical="center", horizontal="center")

def prevent_row_split(row):
    """Prevent a table row from splitting across pages"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    cantSplit.set(qn('w:val'), "true")
    trPr.append(cantSplit)

def convert_to_pdf(input_docx, output_pdf):
    """
    Convert DOCX to PDF using `docx2pdf` on Windows and LibreOffice on Linux.
    """
    system = platform.system()

    if system == "Windows":
        try:
            import pythoncom
            from docx2pdf import convert
            
            pythoncom.CoInitialize()  # Initialize COM for Windows
            convert(input_docx, output_pdf)
            pythoncom.CoUninitialize()  # Clean up COM
            
            return True
        except Exception as e:
            logger.error(f"Error using docx2pdf on Windows: {str(e)}")
            return False

    else:  # Linux/macOS
        try:
            # Find the correct LibreOffice binary
            libreoffice_commands = ["libreoffice", "soffice"]
            command = next((cmd for cmd in libreoffice_commands if subprocess.run([cmd, "--version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False).returncode == 0), None)
            
            if not command:
                logger.error("Neither LibreOffice nor soffice found. Please install LibreOffice.")
                return False

            output_dir = os.path.dirname(output_pdf) or "."

            process = subprocess.run(
                [command, "--headless", "--convert-to", "pdf", "--outdir", output_dir, input_docx],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=False
            )

            if process.returncode != 0:
                logger.error(f"Error converting with LibreOffice: {process.stderr.decode()}")
                return False

            # Rename output file if necessary
            input_basename = os.path.splitext(os.path.basename(input_docx))[0]
            libreoffice_output = os.path.join(output_dir, f"{input_basename}.pdf")

            if libreoffice_output != output_pdf:
                os.rename(libreoffice_output, output_pdf)

            return True

        except Exception as e:
            logger.error(f"Error converting to PDF on Linux/Mac: {str(e)}")
            return False


@app.route('/generate_doc', methods=['POST'])
def generate_document():
    try:
        # Validate request data
        if not request.is_json:
            return jsonify({
                "success": False,
                "error": "Request must be JSON"
            }), 400

        # Get outward number and file data from request
        data = request.json
        if not data:
            return jsonify({
                "success": False,
                "error": "No JSON data received"
            }), 400

        # outward_number = data.get('outwardNumber')
        outward_number = '1069'
        coordinates_data = data.get('fileData')

        if not outward_number or not coordinates_data:
            return jsonify({
                "success": False,
                "error": "Missing required fields: outwardNumber or fileData"
            }), 400

        logger.info(f"Processing outward number: {outward_number}")
        logger.info(f"Coordinates data: {coordinates_data}")

        # Fetch user data from API
        try:
            user_response = requests.get(f'http://localhost:5000/get_user/{outward_number}')
            user_response.raise_for_status()
            user_data = user_response.json()
        except requests.exceptions.RequestException as e:
            logger.error(f"Error fetching user data: {str(e)}")
            return jsonify({
                "success": False,
                "error": f"Failed to fetch user data: {str(e)}"
            }), 500

        # Verify template file exists
        template_path = "MOD 3.docx"
        if not os.path.exists(template_path):
            return jsonify({
                "success": False,
                "error": "Template file not found"
            }), 500

        # Create document
        try:
            docmonarch = Document(template_path)
        except Exception as e:
            logger.error(f"Error creating document: {str(e)}")
            return jsonify({
                "success": False,
                "error": f"Failed to create document: {str(e)}"
            }), 500

        # Update user information
        try:
            name_on_certificate = user_data["user"]["nameoncertificate"]
            corresponding_Address = user_data["user"]["correspondanceadress"]
            Survey_no = f"Survey No:" + user_data["user"]["gutnumber"]
            site_adress = f"Village :{user_data['user']['village']} Taluka :{user_data['user']['taluka']} District :{user_data['user']['district']} Pincode :{user_data['user']['pincode']}"

            # Update date in all headers - Get current date in the desired format (DD/MM/YYYY)
            import datetime
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # Add this import
            current_date = datetime.datetime.now().strftime("%d/%m/%Y")
            logger.info(f"Current Date: {current_date}")  # Log the date
            
            # Update the date in every section's header
            date_updated = False
            for section in docmonarch.sections:
                header = section.header
                
                # First check paragraphs in the header
                for paragraph in header.paragraphs:
                    text = paragraph.text
                    logger.info(f"Header Paragraph Text: '{text}'")  # Log the text content
                    
                    if "Date" in text:
                        paragraph.clear()
                        run = paragraph.add_run(f"Date - {current_date}")
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.font.bold = True
                        logger.info(f"Updated paragraph with date: {paragraph.text}")
                        date_updated = True
                
                # Check for date in header tables
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                text = paragraph.text
                                logger.info(f"Header Table Cell Text: '{text}'")
                                
                                if "Date" in text:
                                    paragraph.clear()
                                    run = paragraph.add_run(f"Date - {current_date}")
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(12)
                                    run.font.bold = True
                                    logger.info(f"Updated table cell with date: {paragraph.text}")
                                    date_updated = True
            
            # If no date field was found in headers, check the main document body
            if not date_updated:
                logger.info("No date field found in headers, checking document body")
                for paragraph in docmonarch.paragraphs:
                    if "Date" in paragraph.text:
                        paragraph.clear()
                        run = paragraph.add_run(f"Date - {current_date}")
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.font.bold = True
                        logger.info(f"Updated body paragraph with date: {paragraph.text}")
                        date_updated = True
                        break
            
            # If still no date field found, try to add it to the header
            if not date_updated:
                logger.info("No date field found, adding to first section header")
                if docmonarch.sections:
                    header = docmonarch.sections[0].header
                    paragraph = header.add_paragraph(f"Date - {current_date}")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.right_indent = Inches(0.5)  # Right padding
                    # paragraph_format.line_spacing = Pt(6)  # Adjust line spacing as needed
                    # paragraph_format.space_after = Pt(10)  # Bottom margin
                    section = docmonarch.sections[0]
                    section.top_margin = Inches(0.5)  # Adjust the top margin of the section


                    

                    run = paragraph.runs[0]
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    # run.font.bold = True
                    logger.info(f"Added new date field to header: {paragraph.text}")

            # Update paragraphs
            if len(docmonarch.paragraphs) > 6:
                paragraph = docmonarch.paragraphs[6]
                paragraph.clear()
                run = paragraph.add_run(name_on_certificate)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
                set_paragraph_format(paragraph)

            if len(docmonarch.paragraphs) > 8:
                paragraph = docmonarch.paragraphs[9]
                paragraph.clear()
                run = paragraph.add_run(corresponding_Address)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
                set_paragraph_format(paragraph)
                paragraph.paragraph_format.left_indent = Inches(0.2)

            if len(docmonarch.paragraphs) > 15:
                paragraph = docmonarch.paragraphs[7]  
                paragraph.clear()
                run = paragraph.add_run(Survey_no)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
            
            if len(docmonarch.paragraphs) > 15:
                paragraph = docmonarch.paragraphs[13]  
                paragraph.clear()
                run = paragraph.add_run(Survey_no)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                # run.font.bold = True

            if len(docmonarch.paragraphs) > 17:
                paragraph = docmonarch.paragraphs[8]  
                paragraph.clear()
                run = paragraph.add_run(site_adress)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True

            if len(docmonarch.paragraphs) > 17:
                paragraph = docmonarch.paragraphs[15]  
                paragraph.clear()
                run = paragraph.add_run(site_adress)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                # run.font.bold = True

            # Update table with coordinates data
            if docmonarch.tables:
                if len(docmonarch.tables) > 1:
                    # for table in docmonarch.tables[1:]:
                    #     table._element.getparent().remove(table._element)
                
                    table = docmonarch.tables[0]
                    table1 = docmonarch.tables[1]
                    row_index = 3

                    # Remove existing data rows
                    for _ in range(len(table.rows) - row_index):
                        table._element.remove(table.rows[row_index]._element)

                    for _ in range(len(table1.rows) - row_index):
                        table1._element.remove(table1.rows[row_index]._element)

                    serial_number = 1
                    for entry in coordinates_data:
                        pattern = r"^\s*[Pp]"
                        if re.match(pattern, entry['P_name']):

                            new_row = table.add_row()
                            prevent_row_split(new_row)

                            for i, cell in enumerate(new_row.cells):
                                for paragraph in cell.paragraphs:
                                    paragraph._element.getparent().remove(paragraph._element)
                                # Add new paragraph with controlled formatting
                                paragraph = cell.add_paragraph()
                                paragraph_format = paragraph.paragraph_format
                                paragraph_format.space_before = Pt(0)  # Remove space before
                                paragraph_format.space_after = Pt(0)   # Remove space after
                                paragraph_format.line_spacing = 1.0    # Single line spacing

                                
                                if i == 0:
                                    run = paragraph.add_run(str(serial_number))
                                elif i == 1:
                                    run = paragraph.add_run(f"Point No. {entry['P_name']} :- Differential GPS Observation taken on Ground IN STATIC mode")
                                elif i == 2:
                                    run = paragraph.add_run(entry['latitude_dms'])
                                elif i == 3:
                                    run = paragraph.add_run(entry['longitude_dms'])
                                elif i == 4:
                                    run = paragraph.add_run(str(entry['Height']))
                                elif i == 5 and 'distances_to_reference_points_km' in entry:
                                    run = paragraph.add_run(f"{entry['distances_to_reference_points_km']['NDA']:.2f} KM")
                                elif i == 6 and 'distances_to_reference_points_km' in entry:
                                    run = paragraph.add_run(f"{entry['distances_to_reference_points_km']['loh']:.2f} KM")
                                elif i == 7 and 'boundary_distances' in entry:
                                    run = paragraph.add_run(f"NDA Min Distance: {entry['boundary_distances']['NDAboundaryMinDistance']:.2f} KM\nLohgaon Min Distance: {entry['boundary_distances']['LohgaonBoundaryMinDistance']:.2f} KM")
                                
                                # Format the run
                                run.font.name = "Arial"
                                run.font.size = Pt(12)

                        else:
                            new_row = table1.add_row()
                            prevent_row_split(new_row)

                            for i, cell in enumerate(new_row.cells):
                                for paragraph in cell.paragraphs:
                                    paragraph._element.getparent().remove(paragraph._element)
                                # Add new paragraph with controlled formatting
                                paragraph = cell.add_paragraph()
                                paragraph_format = paragraph.paragraph_format
                                paragraph_format.space_before = Pt(0)  # Remove space before
                                paragraph_format.space_after = Pt(0)   # Remove space after
                                paragraph_format.line_spacing = 1.0    # Single line spacing

                                
                                if i == 0:
                                    run = paragraph.add_run(str(serial_number))
                                elif i == 1:
                                    run = paragraph.add_run(f"Point No. {entry['P_name']} :- Differential GPS Observation taken on Ground IN STATIC mode")
                                elif i == 2:
                                    run = paragraph.add_run(entry['latitude_dms'])
                                elif i == 3:
                                    run = paragraph.add_run(entry['longitude_dms'])
                                elif i == 4:
                                    run = paragraph.add_run(str(entry['Height']))
                                elif i == 5 and 'distances_to_reference_points_km' in entry:
                                    run = paragraph.add_run(f"{entry['distances_to_reference_points_km']['NDA']:.2f} KM")
                                elif i == 6 and 'distances_to_reference_points_km' in entry:
                                    run = paragraph.add_run(f"{entry['distances_to_reference_points_km']['loh']:.2f} KM")
                                elif i == 7 and 'boundary_distances' in entry:
                                    run = paragraph.add_run(f"NDA Min Distance: {entry['boundary_distances']['NDAboundaryMinDistance']:.2f} KM\nLohgaon Min Distance: {entry['boundary_distances']['LohgaonBoundaryMinDistance']:.2f} KM")
                                
                                # Format the run
                                run.font.name = "Arial"
                                run.font.size = Pt(12)
                        
                        # Set the font and formatting for each cell
                        for cell in new_row.cells:
                            set_cell_alignment(cell, vertical="center", horizontal="center")  # Center both vertically and horizontally

                        serial_number += 1

                    set_table_borders(table)
                    set_table_borders(table1)
                    adjust_table_cell_alignments(table)
                    adjust_table_cell_alignments(table1)
                

            # Save document
            output_docx = 'modified_output.docx'
            output_pdf = 'modified_output.pdf'
            docmonarch.save(output_docx)     
            
            
            docmonarch = None  # Release the document
            
            # Convert to PDF using platform-specific method
            if convert_to_pdf(output_docx, output_pdf):
                if not os.path.exists(output_pdf):
                    raise Exception("PDF file was not created")

                return jsonify({
                    "success": True,
                    "message": "Document generated successfully",
                    "docPath": os.path.abspath(output_docx),
                    "pdfPath": os.path.abspath(output_pdf)
                })
            else:
                return jsonify({
                    "success": False,
                    "error": "Failed to convert document to PDF"
                }), 500

        except Exception as e:
            logger.error(f"Error during document generation: {str(e)}")
            return jsonify({
                "success": False,
                "error": f"Error during document generation: {str(e)}"
            }), 500

    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        return jsonify({
            "success": False,
            "error": f"Unexpected error: {str(e)}"
        }), 500




@app.route('/get-doc')
def get_document():
    try:
        return send_file(
            'modified_output.docx',
            as_attachment=True,
            download_name='modified_output.docx'
        )
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

# API for a PDF viewer

@app.route('/view-pdf')
def view_pdf():
    try:
        return send_file(
            'modified_output.pdf',
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/download-pdf/<outward_number>')
def download_pdf(outward_number):
    try:
        return send_file(
            'modified_output.pdf',
            as_attachment=True,
            download_name=f'{outward_number}.pdf',
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


if __name__ == '__main__':
        app.run(debug=True, host='0.0.0.0', port=5000)















